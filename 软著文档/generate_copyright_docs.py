# -*- coding: utf-8 -*-
"""
软著文档生成脚本（DOCX格式）
生成：
1. Android 源代码文档（前30页+后30页，共60页）
2. 安卓用户使用说明书（18-30页）
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = r"D:\Project\Android\RockiotTag"
SRC_BASE = os.path.join(PROJECT_DIR, "app", "src", "main", "java", "com", "RockiotTag", "tag")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "软著文档")

HEADER_TEXT = "RockiotTag 智能防丢标签管理 APP（Android 版）V1.0"
LINES_PER_PAGE_CODE = 50
TOTAL_PAGES_CODE = 60
FRONT_PAGES = 30
BACK_PAGES = 30

CODE_FILES_ORDERED = [
    "RockiotTagApplication.java",
    "LoginActivity.java",
    "MainActivity.java",
    "viewmodel/MainViewModel.java",
    "helper/MainActivityUIHelper.java",
    "helper/MainActivityDeviceHelper.java",
    "BLEManager.java",
    "bluetooth/OptimizedBLEScanner.java",
    "bluetooth/BluetoothScanResult.java",
    "BLEForegroundService.java",
    "viewmodel/BleViewModel.java",
    "repository/BLERepository.java",
    "util/BLETagFilter.java",
    "util/ProximityDetector.java",
    "AddDeviceActivity.java",
    "viewmodel/AddDeviceViewModel.java",
    "CustomCaptureActivity.java",
    "TagAdapter.java",
    "util/DeviceNicknameFilter.java",
    "UnboundDeviceManager.java",
    "usecase/SaveDeviceUseCase.java",
    "usecase/DeleteDeviceUseCase.java",
    "usecase/SelectDeviceUseCase.java",
    "usecase/SyncDevicesUseCase.java",
    "GeofenceActivity.java",
    "viewmodel/GeofenceViewModel.java",
    "usecase/TriggerBuzzerUseCase.java",
    "DeviceListActivity.java",
    "viewmodel/DeviceListViewModel.java",
    "BoundDeviceAdapter.java",
    "DeviceAdapter.java",
    "adapter/DeviceListAdapter.java",
    "usecase/GetDeviceInfoUseCase.java",
    "usecase/UpdateDeviceLocationUseCase.java",
    "TrackActivity.java",
    "viewmodel/TrackViewModel.java",
    "viewmodel/TrackViewModelFactory.java",
    "usecase/LoadTrackDataUseCase.java",
    "usecase/TrackStatisticsUseCase.java",
    "usecase/DetectStayPointsUseCase.java",
    "fragment/TrackMapFragment.java",
    "fragment/TrackPlaybackFragment.java",
    "fragment/TrackToolbarFragment.java",
    "fragment/TrackDateSelectorFragment.java",
    "helper/TrackMapRenderer.java",
    "helper/TrackPlaybackHelper.java",
    "helper/TrackAddressHelper.java",
    "helper/TrackCameraHelper.java",
    "helper/TrackDataProcessor.java",
    "helper/TrackStatisticsHelper.java",
    "manager/TrackMapController.java",
    "manager/TrackPlaybackManager.java",
    "MapManager.java",
    "map/amap/AMapManager.java",
    "map/amap/AMapLocationService.java",
    "map/amap/AMapGeocoder.java",
    "map/google/GoogleMapManager.java",
    "map/google/GoogleLocationService.java",
    "map/google/GoogleGeocoderService.java",
    "viewmodel/MapViewModel.java",
    "fragment/MapFragment.java",
    "WebViewMapActivity.java",
    "viewmodel/WebViewMapViewModel.java",
    "util/MapMarkerHelper.java",
    "util/GoogleMapMarkerHelper.java",
    "util/GoogleMapTrackRenderer.java",
    "CoordinateUtils.java",
    "Device.java",
    "model/TagDevice.java",
    "model/DeviceLocation.java",
    "model/DeviceTag.java",
    "model/LocationData.java",
    "model/PhoneLocation.java",
    "model/ApiResponse.java",
    "model/Resource.java",
    "model/StayPoint.java",
    "StayPoint.java",
    "LocationRecord.java",
    "repository/DeviceRepository.java",
    "repository/DataRepository.java",
    "repository/LocationRepository.java",
    "room/AppDatabase.java",
    "room/DeviceEntity.java",
    "room/DeviceDao.java",
    "room/LocationRecordEntity.java",
    "room/LocationRecordDao.java",
    "room/AddressCacheEntity.java",
    "room/AddressCacheDao.java",
    "room/Migration_1_2.java",
    "DatabaseHelper.java",
    "NewApiService.java",
    "network/HttpHelper.java",
    "network/NetworkManager.java",
    "integration/LocationOptimizationManager.java",
    "ApiConfig.java",
    "SharedPreferencesManager.java",
    "LanguageUtils.java",
    "CrowdSourcingManager.java",
    "LocationManager.java",
    "GoogleGeocoder.java",
    "fragment/DeviceControlFragment.java",
    "fragment/DeviceInfoFragment.java",
    "fragment/ControlPanelFragment.java",
    "provider/LocationProvider.java",
    "location/PhoneLocationService.java",
    "location/LocationManager.java",
    "data/DeviceDataManager.java",
    "cache/IconCache.java",
    "ui/TimeRefreshManager.java",
    "usecase/BaseUseCase.java",
    "usecase/ReverseGeocodeUseCase.java",
    "util/GlobalExceptionHandler.java",
    "util/MemoryLeakDetector.java",
    "util/SafeExecutor.java",
    "util/SafeHandler.java",
    "util/RetryableTask.java",
    "util/LifecycleResourceManager.java",
    "util/LocationKalmanFilter.java",
    "util/LocationValidator.java",
    "util/TimeFormatter.java",
    "util/TimePickerHelper.java",
    "util/TrackCalculator.java",
    "util/DialogHelper.java",
    "util/ErrorMessageResolver.java",
    "util/AddressCache.java",
    "util/GeocodeHelper.java",
    "util/GoogleGeocodingAPI.java",
    "util/ImageLoader.java",
    "util/DeviceInfoUpdater.java",
    "util/UserFriendlyError.java",
]


def clean_line(text):
    """移除无法渲染的 emoji 和特殊字符"""
    replacements = {
        '\u26a1': '[!]',
        '\u274c': '[X]',
        '\u2713': '[V]',
        '\u26a0': '[!]',
        '\ufe0f': '',
        '\U0001f4f1': '',
        '\U0001f4e2': '',
        '\U0001f517': '',
        '\u2705': '[V]',
        '\u2b55': '[O]',
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    result = []
    for c in text:
        cp = ord(c)
        if cp < 0x2000 or (0x3000 <= cp <= 0x9fff) or (0xff00 <= cp <= 0xffef):
            result.append(c)
        elif 0x2000 <= cp < 0x2100:
            result.append(' ')
        else:
            result.append('')
    return ''.join(result)


def read_effective_lines(filepath):
    lines = []
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                stripped = line.rstrip("\n\r")
                if stripped.strip():
                    lines.append(stripped)
    except FileNotFoundError:
        print(f"  [WARN] 文件不存在，跳过: {filepath}")
    return lines


def collect_all_code_lines():
    all_lines = []
    for rel_path in CODE_FILES_ORDERED:
        filepath = os.path.join(SRC_BASE, rel_path.replace("/", os.sep))
        lines = read_effective_lines(filepath)
        if lines:
            all_lines.append("// ========== 文件: {} ==========".format(rel_path))
            all_lines.extend(lines)
    return all_lines


def add_header(doc, text=HEADER_TEXT):
    """为文档添加页眉"""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    # 清除默认段落
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_page_number(doc):
    """为文档添加页码（右上角）"""
    section = doc.sections[0]
    header = section.header
    # 在页眉段落中添加制表符和页码
    p = header.paragraphs[0]
    # 添加右对齐的页码
    run = p.add_run("\t")
    # 使用域代码插入页码
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2 = p.add_run()
    run2._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run3 = p.add_run()
    run3._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run4 = p.add_run()
    run4._element.append(fldChar2)
    # 右对齐
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_cell_shading(cell, color):
    """设置表格单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def generate_source_code_docx():
    print("\n=== 生成源代码文档 ===")

    effective_lines = collect_all_code_lines()
    print(f"  有效代码总行数: {len(effective_lines)}")

    total_lines_needed = TOTAL_PAGES_CODE * LINES_PER_PAGE_CODE
    if len(effective_lines) < total_lines_needed:
        print(f"  [WARN] 有效代码行数不足: {len(effective_lines)} < {total_lines_needed}")

    front_lines = effective_lines[:FRONT_PAGES * LINES_PER_PAGE_CODE]
    back_lines = effective_lines[-(BACK_PAGES * LINES_PER_PAGE_CODE):]
    selected_lines = front_lines + back_lines
    print(f"  前30页行数: {len(front_lines)}, 后30页行数: {len(back_lines)}")

    doc = Document()

    # 页面设置：A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # 页眉+页码
    add_header(doc)
    add_page_number(doc)

    # 用表格实现代码排版（无边框，单列，每行一个单元格）
    # 这样可以精确控制每页行数
    line_idx = 0
    for page in range(TOTAL_PAGES_CODE):
        if page > 0:
            # 分页符
            doc.add_page_break()

        # 每页50行代码
        for i in range(LINES_PER_PAGE_CODE):
            if line_idx < len(selected_lines):
                line = selected_lines[line_idx]
                line_idx += 1
                line = clean_line(line)
                if len(line) > 95:
                    line = line[:92] + "..."

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(10)

                if line.startswith("// ========== "):
                    run = p.add_run(line)
                    run.font.size = Pt(7)
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 128)
                else:
                    run = p.add_run(line)
                    run.font.size = Pt(7)
                    run.font.name = "Courier New"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                # 空行补齐
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(10)

    output_path = os.path.join(OUTPUT_DIR, "RockiotTag_源代码.docx")
    doc.save(output_path)
    print(f"  源代码文档已生成: {output_path}")
    print(f"  总页数: {TOTAL_PAGES_CODE}")


# ============================================================
# 说明书辅助函数
# ============================================================

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(doc, text, font_size=11, bold=False, align=None, font_name="宋体", first_indent=None):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.bold = bold
    return p


def add_lines(doc, lines, font_size=11, font_name="宋体"):
    """添加多行文本"""
    for line in lines:
        add_para(doc, line, font_size=font_size, font_name=font_name)


def add_screenshot_placeholder(doc, title):
    """添加截图占位符"""
    add_para(doc, "【" + title + "】", font_size=10, bold=True)
    # 添加一个带边框的占位框
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n（请在此处插入安卓手机" + title + "截图）\n\n")
    run.font.size = Pt(10)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(128, 128, 128)
    # 设置单元格高度
    cell.height = Cm(5)
    doc.add_paragraph()  # 间距


def generate_manual_docx():
    print("\n=== 生成用户使用说明书 ===")

    doc = Document()

    # 页面设置：A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 页眉+页码
    add_header(doc)
    add_page_number(doc)

    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph()

    add_para(doc, "RockiotTag", font_size=30, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_name="黑体")
    add_para(doc, "智能防丢标签管理 APP", font_size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_name="黑体")
    doc.add_paragraph()
    add_para(doc, "安卓用户使用说明书", font_size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_name="黑体")
    add_para(doc, "（Android 版）V1.0", font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, "运行环境：Android 8.0 及以上", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "开发语言：Java", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "版本号：V1.0", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "版权所有 © RockiotTag", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ========== 目录 ==========
    doc.add_page_break()
    add_heading_styled(doc, "目  录", level=1)

    toc_items = [
        "一、软件概述",
        "    1.1 产品简介",
        "    1.2 运行环境",
        "    1.3 开发语言与技术架构",
        "    1.4 产品用途与适用场景",
        "二、软件安装与启动",
        "    2.1 安装方式",
        "    2.2 启动应用",
        "    2.3 权限说明",
        "三、功能说明与操作指南",
        "    3.1 启动页与登录",
        "    3.2 首页地图与设备概览",
        "    3.3 添加标签设备",
        "    3.4 蓝牙连接与配对",
        "    3.5 设备列表管理",
        "    3.6 防丢设置与告警",
        "    3.7 电子围栏",
        "    3.8 轨迹追踪与回放",
        "    3.9 蜂鸣器寻物",
        "    3.10 设备信息与设置",
        "四、常见问题与故障排除",
        "五、技术支持与联系方式",
    ]
    for item in toc_items:
        add_para(doc, item, font_size=12)

    # ========== 一、软件概述 ==========
    doc.add_page_break()
    add_heading_styled(doc, "一、软件概述", level=1)

    add_heading_styled(doc, "1.1 产品简介", level=2)
    add_lines(doc, [
        "RockiotTag 智能防丢标签管理 APP 是一款专为 Android 智能手机用户设计的蓝牙防丢器管理应用。该应用通过低功耗蓝牙（BLE）技术，与 RockiotTag 智能防丢标签硬件设备进行无线连接，实现对贵重物品的实时定位、防丢告警、轨迹追踪等功能。用户可将防丢标签挂载在钥匙、钱包、行李箱、宠物项圈等物品上，通过 APP 即可随时掌握物品位置，有效防止物品丢失。",
        "",
        "本应用采用 MVVM 架构设计，集成了高德地图和谷歌地图双地图引擎，支持全球定位服务。同时应用内置了 BLE 距离检测、电子围栏、轨迹回放、蜂鸣器寻物等多种实用功能，为用户提供全方位的防丢保护方案。应用界面简洁直观，操作便捷，适合各年龄段用户使用。",
        "",
        "RockiotTag APP 的核心价值在于：通过蓝牙近场检测技术，在用户与贵重物品之间建立一道无形的防护线。当物品离开安全距离时，APP 会立即发出告警通知，帮助用户及时发现并找回物品，从根本上解决物品丢失的困扰。",
    ])

    add_heading_styled(doc, "1.2 运行环境", level=2)
    add_lines(doc, [
        "本应用的运行环境要求如下：",
        "",
        "  操作系统：Android 8.0（API Level 26）及以上版本",
        "  硬件要求：支持蓝牙 4.0 低功耗（BLE）的 Android 智能手机",
        "  网络要求：需要移动网络或 Wi-Fi 连接（用于地图加载和云端数据同步）",
        "  定位服务：需开启手机 GPS 定位功能",
        "  蓝牙权限：需开启手机蓝牙功能并授予相关权限",
        "  存储空间：应用安装包约 30MB，运行时需约 100MB 存储空间",
        "  屏幕适配：支持 4.7 英寸及以上屏幕尺寸的 Android 设备",
        "  推荐配置：Android 10.0 及以上，4GB RAM 及以上",
        "",
        "注意：本应用仅支持 Android 操作系统，不支持 iOS 系统。如需在 iPhone 上使用，请下载 iOS 版本的应用。",
    ])

    add_heading_styled(doc, "1.3 开发语言与技术架构", level=2)
    add_lines(doc, [
        "本应用的技术架构信息如下：",
        "",
        "  开发语言：Java",
        "  最低 SDK 版本：Android 5.0（API Level 21）",
        "  目标 SDK 版本：Android 14（API Level 35）",
        "  架构模式：MVVM（Model-View-ViewModel）",
        "  数据存储：Room 数据库 + SharedPreferences",
        "  网络通信：Volley + 自定义 HTTP 请求",
        "  地图服务：高德地图 SDK + Google Maps SDK（双引擎切换）",
        "  蓝牙通信：Android BLE API + 自定义扫描优化策略",
        "  扫码功能：ML Kit 条码扫描 + CameraX",
        "  图片加载：Glide 图片加载框架",
        "  后台服务：前台服务保障后台 BLE 扫描持续运行",
        "",
        "应用采用 MVVM 架构模式，将界面展示（View）、业务逻辑（ViewModel）和数据层（Model）清晰分离，确保代码的可维护性和可测试性。数据层使用 Room 数据库进行本地持久化存储，同时通过网络请求与云端服务器进行数据同步。",
    ])

    add_heading_styled(doc, "1.4 产品用途与适用场景", level=2)
    add_lines(doc, [
        "RockiotTag 智能防丢标签管理 APP 主要用于以下场景：",
        "",
        "  1. 钥匙防丢：将标签挂在钥匙扣上，离开一定距离自动告警，再也不用担心找不到钥匙",
        "  2. 钱包防丢：将标签放入钱包内，实时监控钱包位置，防止钱包遗失",
        "  3. 行李追踪：旅行时将标签放入行李箱，追踪行李位置，防止行李丢失或被错拿",
        "  4. 宠物定位：将标签挂在宠物项圈上，防止宠物走失，随时掌握宠物位置",
        "  5. 车辆定位：将标签放在车内，记录车辆停放位置，方便在大型停车场找车",
        "  6. 其他贵重物品：任何需要防丢追踪的物品均可使用，如笔记本电脑包、相机包等",
        "",
        "本产品通过蓝牙近场检测和云端定位双重保障，结合电子围栏和轨迹追踪功能，为用户的贵重物品提供全方位的防丢保护。无论是日常生活中的小物件防丢，还是旅行途中的行李追踪，RockiotTag 都能提供可靠的保障。",
    ])

    # ========== 二、软件安装与启动 ==========
    doc.add_page_break()
    add_heading_styled(doc, "二、软件安装与启动", level=1)

    add_heading_styled(doc, "2.1 安装方式", level=2)
    add_lines(doc, [
        "用户可通过以下方式安装 RockiotTag APP：",
        "",
        "方式一：Google Play 商店安装（推荐）",
        "  1. 在 Android 手机上打开 Google Play 商店应用",
        "  2. 在搜索栏中输入『RockiotTag』进行搜索",
        "  3. 找到『RockiotTag 智能防丢标签管理』应用",
        "  4. 点击『安装』按钮，等待下载和安装完成",
        "  5. 安装完成后，点击『打开』即可启动应用",
        "",
        "方式二：APK 安装包安装",
        "  1. 从官方网站下载 RockiotTag APK 安装包",
        "  2. 在手机设置中开启『允许安装未知来源应用』选项",
        "  3. 打开下载的 APK 文件，按照提示完成安装",
        "  4. 安装完成后，在桌面找到 RockiotTag 图标点击启动",
        "",
        "建议优先使用 Google Play 商店安装，以确保获得最新版本和安全更新。",
    ])

    add_heading_styled(doc, "2.2 启动应用", level=2)
    add_lines(doc, [
        "安装完成后，可通过以下方式启动应用：",
        "",
        "  1. 在手机桌面找到 RockiotTag 应用图标，点击启动",
        "  2. 应用启动后，首先显示启动页（Splash Screen），展示品牌标识",
        "  3. 启动页加载完成后，自动跳转至应用首页",
        "  4. 首次使用时，应用会引导用户授予必要权限",
        "  5. 权限授予后，即可正常使用应用的所有功能",
        "",
        "启动页加载过程包括：初始化应用组件、检查蓝牙状态、加载本地数据库等操作。通常在2秒内完成加载。如果首次启动时间较长，属于正常现象，后续启动会明显加快。",
    ])
    add_screenshot_placeholder(doc, "启动页界面")
    add_screenshot_placeholder(doc, "权限申请界面")

    add_heading_styled(doc, "2.3 权限说明", level=2)
    add_lines(doc, [
        "安装过程中及首次启动时，应用会请求以下权限：",
        "",
        "  蓝牙权限：用于扫描和连接防丢标签设备，是应用核心功能的基础",
        "  位置权限：用于地图定位和距离计算，同时 BLE 扫描也需要位置权限",
        "  相机权限：用于扫描标签设备上的二维码，快速添加设备",
        "  存储权限：用于缓存地图数据和设备信息，提升应用响应速度",
        "  通知权限：用于发送防丢告警通知，确保用户及时收到告警信息",
        "  后台服务权限：用于后台持续监控设备连接状态，保障防丢功能正常运行",
        "",
        "以上权限均为应用核心功能所必需，应用不会收集或上传用户的隐私数据。位置信息仅用于本地地图展示和距离计算，不会发送至第三方服务器。",
    ])

    # ========== 三、功能说明与操作指南 ==========
    doc.add_page_break()
    add_heading_styled(doc, "三、功能说明与操作指南", level=1)

    add_heading_styled(doc, "3.1 启动页与登录", level=2)
    add_lines(doc, [
        "应用启动后，首先展示启动页界面，显示 RockiotTag 品牌标识和加载进度。启动页加载完成后，系统自动跳转至应用主界面（首页）。应用无需注册登录即可使用基本功能，设备数据通过本地数据库存储，同时支持云端同步功能。",
        "",
        "操作步骤：",
        "  1. 点击桌面 RockiotTag 图标启动应用",
        "  2. 等待启动页加载完成（约 2 秒）",
        "  3. 首次启动时，按照提示授予蓝牙、位置等必要权限",
        "  4. 权限授予后，自动进入应用首页",
        "  5. 后续启动无需重复授权，直接进入首页",
    ])
    add_screenshot_placeholder(doc, "启动页界面")

    add_heading_styled(doc, "3.2 首页地图与设备概览", level=2)
    add_lines(doc, [
        "应用首页为地图视图，展示用户当前所在位置及已绑定设备的位置信息。首页顶部为地图区域，底部为设备信息面板，可上下滑动查看更多内容。",
        "",
        "首页功能区域说明：",
        "  - 地图区域：显示当前定位和设备位置标记，支持缩放和拖拽操作",
        "  - 设备信息卡片：显示当前选中设备的名称、电量、连接状态、距离等信息",
        "  - 距离指示：显示手机与标签设备的估算距离，基于 BLE 信号强度计算",
        "  - 蜂鸣器按钮：点击触发标签蜂鸣器，方便快速寻物",
        "  - 导航栏：底部导航切换不同功能页面（首页、设备、设置）",
        "",
        "地图支持高德地图和谷歌地图双引擎，根据用户所在地区自动切换。国内默认使用高德地图，海外默认使用谷歌地图，确保全球范围内均可正常使用定位服务。用户也可在设置中手动切换地图引擎。",
    ])
    add_screenshot_placeholder(doc, "首页地图界面")
    add_screenshot_placeholder(doc, "设备信息面板")

    add_heading_styled(doc, "3.3 添加标签设备", level=2)
    add_lines(doc, [
        "用户可通过添加设备功能，将新的 RockiotTag 防丢标签绑定到 APP 中进行管理。添加设备支持扫码添加和手动输入设备号两种方式。",
        "",
        "操作步骤：",
        "  1. 在首页或设备列表页面，点击『+』或『添加设备』按钮",
        "  2. 进入添加设备页面，选择添加方式：",
        "     a. 扫码添加：点击扫码按钮，使用手机摄像头扫描标签上的二维码",
        "     b. 手动输入：在输入框中手动输入设备编号",
        "  3. 选择标签类型（如钥匙标签、钱包标签等）",
        "  4. 为设备设置昵称（如『我的钥匙』、『钱包标签』等）",
        "  5. 点击『确认绑定』按钮完成添加",
        "  6. 绑定成功后，设备将出现在设备列表中",
        "",
        "扫码功能使用 ML Kit 条码识别技术，支持快速扫描二维码。同时应用支持 CameraX 相机预览，提供流畅的扫码体验。在光线不足的环境下，应用会自动开启闪光灯辅助扫描。",
    ])
    add_screenshot_placeholder(doc, "添加设备页面")
    add_screenshot_placeholder(doc, "扫码界面")
    add_screenshot_placeholder(doc, "设备绑定成功界面")

    add_heading_styled(doc, "3.4 蓝牙连接与配对", level=2)
    add_lines(doc, [
        "RockiotTag APP 通过低功耗蓝牙（BLE）与标签设备建立连接，实现数据传输和状态监控。应用内置了优化的 BLE 扫描策略，采用周期扫描与休息交替的方式，在保证连接稳定性的同时最大限度降低手机功耗。",
        "",
        "蓝牙连接流程：",
        "  1. 确保手机蓝牙已开启，并已授予蓝牙权限",
        "  2. 确保标签设备电量充足且处于开启状态",
        "  3. APP 自动扫描附近的 RockiotTag 设备",
        "  4. 扫描到目标设备后，自动建立 BLE 连接",
        "  5. 连接成功后，设备状态显示为『已连接』",
        "  6. 连接状态下可读取设备电量、信号强度等信息",
        "",
        "蓝牙连接特性：",
        "  - 自动重连：设备断开后，APP 会自动尝试重新连接",
        "  - 后台扫描：应用支持后台持续扫描，即使切换到其他应用也能保持监控",
        "  - 多设备管理：支持同时连接和管理多个标签设备",
        "  - 信号强度指示：实时显示 RSSI 信号强度，估算设备距离",
        "  - 前台服务：通过前台服务保障后台 BLE 扫描不被系统杀死",
    ])
    add_screenshot_placeholder(doc, "蓝牙连接界面")
    add_screenshot_placeholder(doc, "设备信号强度界面")

    add_heading_styled(doc, "3.5 设备列表管理", level=2)
    add_lines(doc, [
        "设备列表页面展示用户所有已绑定的标签设备，支持查看设备详情、编辑设备信息、解绑设备等操作。列表采用 RecyclerView + DiffUtil 优化，确保大量设备时的流畅滚动体验。",
        "",
        "设备列表功能：",
        "  - 设备概览：显示每个设备的名称、类型、电量、连接状态",
        "  - 设备详情：点击设备卡片进入详情页，查看位置历史和设备参数",
        "  - 编辑设备：修改设备昵称、标签类型等信息",
        "  - 解绑设备：移除不再使用的标签设备",
        "  - 批量操作：支持多选设备进行批量管理",
        "  - 设备排序：按名称、电量、最近连接时间等排序",
        "",
        "操作步骤：",
        "  1. 点击底部导航栏『设备』进入设备列表页面",
        "  2. 浏览所有已绑定设备，查看各设备状态",
        "  3. 点击设备卡片进入设备详情页",
        "  4. 在详情页可编辑设备信息或进行其他操作",
        "  5. 长按设备卡片可进入多选模式，进行批量操作",
    ])
    add_screenshot_placeholder(doc, "设备列表界面")
    add_screenshot_placeholder(doc, "设备详情界面")

    add_heading_styled(doc, "3.6 防丢设置与告警", level=2)
    add_lines(doc, [
        "防丢告警是 RockiotTag APP 的核心功能。当手机与标签设备的距离超过设定阈值时，APP 会自动发送告警通知，提醒用户注意物品安全。应用通过 BLE 信号强度（RSSI）计算设备距离，并结合卡尔曼滤波算法提高距离估算精度。",
        "",
        "防丢设置选项：",
        "  - 告警距离：设置触发告警的距离阈值（近/中/远三档）",
        "  - 告警方式：选择通知栏提醒、震动、铃声等告警方式",
        "  - 防丢开关：单独控制每个设备的防丢功能开关",
        "  - 断连告警：蓝牙断开连接时自动告警",
        "  - 静音时段：设置特定时段不触发告警",
        "",
        "告警触发流程：",
        "  1. APP 持续监测与标签设备的 BLE 连接状态和信号强度",
        "  2. 当信号强度低于阈值（距离过远）时，触发告警",
        "  3. 告警通知显示在手机通知栏，包含设备名称和告警类型",
        "  4. 用户点击通知可快速跳转至对应设备页面",
        "  5. 设备重新进入范围后，告警自动解除",
    ])
    add_screenshot_placeholder(doc, "防丢设置界面")
    add_screenshot_placeholder(doc, "防丢告警通知")

    add_heading_styled(doc, "3.7 电子围栏", level=2)
    add_lines(doc, [
        "电子围栏功能允许用户在地图上设定一个安全区域，当标签设备离开该区域时，APP 自动发送越界告警通知。用户可自定义围栏的中心点和半径，灵活设置安全范围。",
        "",
        "电子围栏设置步骤：",
        "  1. 在设备详情页，点击『电子围栏』进入设置页面",
        "  2. 地图上显示当前设备位置，默认以设备位置为围栏中心",
        "  3. 拖动地图调整围栏中心点位置",
        "  4. 滑动半径调节条，设置围栏半径（50米-2000米）",
        "  5. 地图上实时显示围栏范围（圆形区域）",
        "  6. 点击『保存』按钮确认围栏设置",
        "  7. 围栏生效后，设备离开围栏范围将触发越界告警",
        "",
        "电子围栏特性：",
        "  - 可视化设置：在地图上直观显示围栏范围",
        "  - 灵活调整：支持拖动和缩放调整围栏位置和大小",
        "  - 多围栏支持：可为同一设备设置多个围栏区域",
        "  - 即时告警：设备越界后立即发送通知",
    ])
    add_screenshot_placeholder(doc, "电子围栏设置界面")
    add_screenshot_placeholder(doc, "电子围栏越界告警")

    add_heading_styled(doc, "3.8 轨迹追踪与回放", level=2)
    add_lines(doc, [
        "轨迹追踪功能记录标签设备的历史位置信息，并在地图上以轨迹线形式展示。用户可查看设备在指定日期内的移动轨迹，支持轨迹回放动画播放。",
        "",
        "轨迹查看操作：",
        "  1. 在设备详情页，点击『轨迹追踪』进入轨迹页面",
        "  2. 页面顶部显示日期选择器，选择要查看的日期",
        "  3. 地图上显示该日期内设备的移动轨迹线",
        "  4. 轨迹线上标注关键位置点和停留点",
        "  5. 底部工具栏提供回放控制：播放/暂停、速度调节",
        "  6. 点击播放按钮，轨迹以动画形式回放设备移动过程",
        "",
        "轨迹功能特性：",
        "  - 日期筛选：选择任意日期查看历史轨迹",
        "  - 停留点检测：自动识别并标注设备停留时间较长的位置",
        "  - 轨迹统计：显示总距离、总时长、平均速度等统计数据",
        "  - 回放动画：支持 1x/2x/4x/8x 多倍速回放",
        "  - 轨迹导出：支持将轨迹数据导出分享",
    ])
    add_screenshot_placeholder(doc, "轨迹追踪界面")
    add_screenshot_placeholder(doc, "轨迹回放界面")

    add_heading_styled(doc, "3.9 蜂鸣器寻物", level=2)
    add_lines(doc, [
        "蜂鸣器寻物功能允许用户通过 APP 远程触发标签设备的蜂鸣器，使标签发出声音，帮助用户快速找到附近的物品。该功能在设备处于蓝牙连接范围内时可用。",
        "",
        "蜂鸣器寻物操作：",
        "  1. 在首页设备信息面板或设备详情页，点击『蜂鸣器』按钮",
        "  2. APP 通过 BLE 向标签设备发送蜂鸣指令",
        "  3. 标签设备收到指令后开始发出蜂鸣声",
        "  4. 根据声音方向和大小，定位标签所在位置",
        "  5. 找到物品后，再次点击按钮停止蜂鸣",
        "",
        "注意事项：",
        "  - 蜂鸣器功能需要设备处于蓝牙连接状态",
        "  - 蜂鸣器有效距离与蓝牙连接范围一致（约 10-30 米）",
        "  - 长时间使用蜂鸣器会增加设备电量消耗",
        "  - 设备电量过低时蜂鸣器可能无法正常工作",
    ])
    add_screenshot_placeholder(doc, "蜂鸣器寻物界面")

    add_heading_styled(doc, "3.10 设备信息与设置", level=2)
    add_lines(doc, [
        "在设备详情页，用户可以查看设备的详细信息并进行相关设置：",
        "",
        "设备信息：",
        "  - 设备名称：显示当前设备的昵称，可点击编辑修改",
        "  - 设备编号：显示设备的唯一标识号",
        "  - 设备类型：显示标签类型（钥匙标签、钱包标签等）",
        "  - 电量信息：显示设备当前电量百分比",
        "  - 固件版本：显示设备固件版本号",
        "  - 连接状态：显示当前蓝牙连接状态",
        "  - 最后定位：显示设备最后一次定位的时间和位置",
        "",
        "设备设置：",
        "  - 防丢开关：开启或关闭该设备的防丢功能",
        "  - 告警距离：调整防丢告警的距离阈值",
        "  - 电子围栏：设置设备的安全区域",
        "  - 蜂鸣器：远程触发设备蜂鸣器",
        "  - 轨迹追踪：查看设备历史轨迹",
        "  - 解绑设备：将该设备从账号中移除",
    ])
    add_screenshot_placeholder(doc, "设备信息界面")
    add_screenshot_placeholder(doc, "设备设置界面")

    # ========== 四、常见问题与故障排除 ==========
    doc.add_page_break()
    add_heading_styled(doc, "四、常见问题与故障排除", level=1)

    faq_items = [
        ("Q1：无法扫描到标签设备怎么办？",
         "A：请检查以下几点：1) 确认手机蓝牙已开启；2) 确认标签设备电量充足且已开机；3) 确认标签设备在蓝牙有效范围内（10米以内）；4) 尝试关闭蓝牙后重新开启；5) 重启 APP 后再次尝试扫描。如果仍然无法扫描，请检查标签设备是否正常工作。"),
        ("Q2：设备经常断开连接怎么办？",
         "A：BLE 连接受环境因素影响较大。请尝试：1) 确保手机与标签之间无障碍物遮挡；2) 避免在 Wi-Fi 信号密集区域使用；3) 检查手机省电模式是否限制了后台蓝牙；4) 将 APP 加入手机省电白名单；5) 确保标签设备电量充足。"),
        ("Q3：地图无法显示或定位不准怎么办？",
         "A：请检查：1) 确认手机 GPS 已开启；2) 确认已授予位置权限；3) 在室内时 GPS 信号较弱，建议到室外开阔区域使用；4) 检查网络连接是否正常（地图加载需要网络）；5) 尝试重启应用。"),
        ("Q4：防丢告警不触发怎么办？",
         "A：请检查：1) 确认防丢功能已开启；2) 确认告警距离设置合理；3) 确认 APP 后台运行权限已授予；4) 检查是否开启了静音时段；5) 确认通知权限已授予；6) 确认设备蓝牙连接正常。"),
        ("Q5：如何更换绑定的手机？",
         "A：在新手机上安装 RockiotTag APP，使用同一账号登录后，设备数据将自动同步。如无需账号，可在旧手机上解绑设备后，在新手机上重新添加绑定。"),
        ("Q6：标签设备电池能用多久？",
         "A：RockiotTag 标签设备采用 CR2032 纽扣电池，正常使用情况下可续航约 6-12 个月。频繁使用蜂鸣器和高频定位会缩短电池寿命。APP 中可查看设备剩余电量。"),
        ("Q7：应用在后台被系统杀死怎么办？",
         "A：Android 系统为节省电量可能会清理后台应用。建议：1) 将 APP 加入电池优化白名单；2) 在系统设置中允许 APP 后台运行；3) 开启 APP 的前台服务通知，降低被杀概率；4) 锁定 APP 在最近任务列表中，防止被一键清理。"),
        ("Q8：电子围栏告警延迟怎么处理？",
         "A：电子围栏告警依赖设备位置上报频率。如果设备通过 BLE 连接，位置更新较实时；如果设备处于离线状态，需等待设备下次上报位置时才能检测越界。建议保持 BLE 连接以获得最佳告警时效。"),
        ("Q9：如何查看设备的历史轨迹？",
         "A：在设备详情页点击『轨迹追踪』，选择要查看的日期即可查看历史轨迹。轨迹功能会自动记录设备的位置变化，支持回放动画和统计数据查看。"),
        ("Q10：应用支持同时连接多少个设备？",
         "A：应用支持同时连接和管理多个标签设备，具体数量取决于手机蓝牙芯片的能力。一般 Android 手机可同时连接 4-7 个 BLE 设备。应用会自动管理连接优先级，确保最常用的设备保持连接状态。"),
    ]

    for q, a in faq_items:
        add_para(doc, q, font_size=11, bold=True)
        add_para(doc, a, font_size=11)

    # ========== 五、技术支持与联系方式 ==========
    doc.add_page_break()
    add_heading_styled(doc, "五、技术支持与联系方式", level=1)
    add_lines(doc, [
        "如在使用过程中遇到任何问题，可通过以下方式获取技术支持：",
        "",
        "1. 在线帮助：应用内『设置-帮助与反馈』页面",
        "2. 电子邮件：发送问题描述至官方客服邮箱",
        "3. 官方网站：访问产品官网获取最新使用指南和常见问题解答",
        "4. 应用更新：定期检查 Google Play 商店更新，获取最新功能优化和问题修复",
        "",
        "使用建议：",
        "  - 定期检查并更新 APP 至最新版本，以获得最佳使用体验",
        "  - 保持标签设备固件为最新版本",
        "  - 建议将 APP 加入手机省电白名单，确保后台监控功能正常运行",
        "  - 长时间不使用时，建议关闭标签设备以节省电量",
        "  - 如遇定位精度问题，建议在开阔区域使用以获得更准确的 GPS 定位",
        "  - 定期检查标签设备电量，及时更换电池以确保设备正常工作",
        "  - 建议为每个标签设置易于识别的昵称，方便管理多个设备",
        "",
        "安全提示：",
        "  - 请勿将标签设备暴露在极端温度或潮湿环境中",
        "  - 标签设备不防水，请避免接触水源",
        "  - 请勿尝试自行拆解或改装标签设备",
        "  - 电池更换时请注意正负极方向",
        "  - 废旧电池请按照当地法规妥善处理",
        "",
        "免责声明：",
        "  本应用提供的定位和距离信息仅供参考，受蓝牙信号、GPS 精度、环境因素等影响，可能存在一定误差。请勿将本应用作为唯一的安全保障手段，对于因定位误差导致的任何损失，开发者不承担任何责任。用户应妥善保管贵重物品，本应用仅作为辅助防丢工具使用。",
        "",
        "",
        "版权所有 © RockiotTag。保留所有权利。",
        "RockiotTag 智能防丢标签管理 APP（Android 版）V1.0",
    ])

    output_path = os.path.join(OUTPUT_DIR, "RockiotTag_用户使用说明书.docx")
    doc.save(output_path)
    print(f"  用户使用说明书已生成: {output_path}")


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_source_code_docx()
    generate_manual_docx()
    print("\n=== 文档生成完成 ===")
    print(f"输出目录: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
