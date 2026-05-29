#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 转换为 Word 文档
使用方法：python convert_to_word.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def read_markdown(filepath):
    """读取 Markdown 文件"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def markdown_to_word(md_content, output_path):
    """将 Markdown 转换为 Word"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 标题处理
        if line.startswith('# '):
            # 一级标题
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        elif line.startswith('## '):
            # 二级标题
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
        elif line.startswith('### '):
            # 三级标题
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
        elif line.startswith('#### '):
            # 四级标题
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
        elif line.startswith('---'):
            # 分隔线
            doc.add_paragraph('─' * 50)
            i += 1
        elif line.startswith('> '):
            # 引用
            p = doc.add_paragraph(line[2:].strip())
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
        elif line.startswith('|'):
            # 表格处理
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if table_lines:
                # 解析表格
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                data_rows = []
                for row_line in table_lines[2:]:  # 跳过分隔行
                    if '|---' not in row_line:
                        row = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        data_rows.append(row)
                
                # 创建表格
                table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
                table.style = 'Table Grid'
                
                # 填充表头
                for j, header in enumerate(headers):
                    table.cell(0, j).text = header
                    table.cell(0, j).paragraphs[0].runs[0].bold = True
                
                # 填充数据
                for row_idx, row_data in enumerate(data_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.cell(row_idx + 1, col_idx).text = cell_data
                
                doc.add_paragraph()  # 表格后添加空行
        elif line.startswith('```'):
            # 代码块处理
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
                p.paragraph_format.left_indent = Inches(0.3)
                i += 1
        elif re.match(r'^\*\*.*\*\*:', line):
            # 加粗文本后跟冒号（字段说明等）
            match = re.match(r'^\*\*(.+?)\*\*:\s*(.*)', line)
            if match:
                p = doc.add_paragraph()
                p.add_run(match.group(1) + ': ').bold = True
                p.add_run(match.group(2))
                i += 1
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            i += 1
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            # 数字列表
            p = doc.add_paragraph(line[3:].strip(), style='List Number')
            i += 1
        else:
            # 普通文本
            if line:
                p = doc.add_paragraph(line)
            i += 1
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ 成功转换为: {output_path}")

if __name__ == '__main__':
    md_file = 'RockiotTag_API_Document.md'
    word_file = 'RockiotTag_API_Document.docx'
    
    print("开始转换...")
    md_content = read_markdown(md_file)
    markdown_to_word(md_content, word_file)
    print("转换完成！")
