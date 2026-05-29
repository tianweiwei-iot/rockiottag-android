"""
Markdown 转 Word 文档转换脚本
将 API_Documentation_DR.md 转换为 API_Documentation_DR.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_run_font(run, font_name='Microsoft YaHei', font_size=10, bold=False, color=None):
    """设置字体样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    # 设置浅灰色背景（通过段落底纹）
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # 使用等宽字体
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)  # 深蓝色

def convert_md_to_docx(md_file, docx_file):
    """将 Markdown 文件转换为 Word 文档"""
    
    # 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)
    
    # 解析 Markdown
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 标题处理
        if line.startswith('# '):
            # 一级标题
            heading = doc.add_heading(line[2:].strip(), level=1)
            for run in heading.runs:
                set_run_font(run, font_size=16, bold=True, color=RGBColor(0, 51, 102))
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            
        elif line.startswith('## '):
            # 二级标题
            heading = doc.add_heading(line[3:].strip(), level=2)
            for run in heading.runs:
                set_run_font(run, font_size=14, bold=True, color=RGBColor(0, 102, 204))
            i += 1
            
        elif line.startswith('### '):
            # 三级标题
            heading = doc.add_heading(line[4:].strip(), level=3)
            for run in heading.runs:
                set_run_font(run, font_size=12, bold=True, color=RGBColor(51, 51, 51))
            i += 1
            
        # 代码块处理
        elif line.startswith('```'):
            # 收集代码块内容
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            add_code_block(doc, code_text)
            i += 1
            
        # 表格处理（简化版）
        elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|---'):
            # 解析表格
            table_data = []
            header_line = line
            i += 2  # 跳过分隔线
            
            # 收集表格数据
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                table_data.append(row)
                i += 1
            
            # 创建表格
            if table_data:
                # 解析表头
                headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
                
                table = doc.add_table(rows=len(table_data) + 1, cols=len(headers))
                table.style = 'Table Grid'
                
                # 填充表头
                for col_idx, header in enumerate(headers):
                    cell = table.cell(0, col_idx)
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run, font_size=9, bold=True)
                
                # 填充数据
                for row_idx, row_data in enumerate(table_data, 1):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(headers):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_data
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    set_run_font(run, font_size=9)
            
        # 列表项处理
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # 处理粗体
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            
        # 普通段落
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"✅ 转换完成！")
    print(f"📄 输出文件：{docx_file}")

def add_formatted_text(paragraph, text):
    """添加带格式的文本（处理粗体、代码等）"""
    # 处理行内代码 `code`
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # 行内代码
            code = part[1:-1]
            run = paragraph.add_run(code)
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(199, 37, 78)  # 红色
        elif '**' in part:
            # 处理粗体 **text**
            sub_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('**') and sub_part.endswith('**'):
                    run = paragraph.add_run(sub_part[2:-2])
                    set_run_font(run, bold=True)
                else:
                    run = paragraph.add_run(sub_part)
                    set_run_font(run)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)

if __name__ == '__main__':
    # 文件路径
    md_file = r'C:\Users\Administrator\Desktop\API_Documentation_DR.md'
    docx_file = r'C:\Users\Administrator\Desktop\API_Documentation_DR.docx'
    
    print("🔄 开始转换...")
    print(f"📝 输入文件：{md_file}")
    print(f"📄 输出文件：{docx_file}")
    print("-" * 50)
    
    try:
        convert_md_to_docx(md_file, docx_file)
        print("-" * 50)
        print("✨ 转换成功！您可以打开 Word 文档查看结果。")
    except Exception as e:
        print(f"❌ 转换失败：{str(e)}")
        import traceback
        traceback.print_exc()
